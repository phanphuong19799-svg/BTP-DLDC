import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==============================
# INLINE FORMAT
# ==============================
def parse_inline_format(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue

        clean = part.strip('*`')
        run = paragraph.add_run(clean)

        if part.startswith('**'):
            run.bold = True
        elif part.startswith('*'):
            run.italic = True
        elif part.startswith('`'):
            run.font.name = 'Courier New'


# ==============================
# TABLE
# ==============================
def add_table_to_doc(doc, table_data):
    if not table_data:
        return

    cols = max(len(r) for r in table_data)
    table = doc.add_table(rows=len(table_data), cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j in range(cols):
            cell = row.cells[j]
            text = row_data[j] if j < len(row_data) else ''
            
            # Xử lý format (in đậm/in nghiêng) trong cell
            cell.text = ''
            if i == 0:
                p = cell.paragraphs[0]
                parse_inline_format(p, text)
                for run in p.runs:
                    run.bold = True
            else:
                parse_inline_format(cell.paragraphs[0], text)


# ==============================
# IMAGE
# ==============================
def handle_image(doc, line, md_path):
    match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
    if match:
        alt, path = match.groups()
        
        # Xử lý đường dẫn tương đối (tính từ vị trí file MD)
        if not os.path.isabs(path):
            md_dir = os.path.dirname(md_path)
            full_path = os.path.abspath(os.path.join(md_dir, path))
        else:
            full_path = path

        if os.path.exists(full_path):
            doc.add_picture(full_path, width=Inches(5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[Missing Image: {full_path}]").alignment = WD_ALIGN_PARAGRAPH.CENTER


# ==============================
# CONVERT MD -> DOCX
# ==============================
def convert_md_to_docx(md_path, docx_path=None, doc=None):
    save_doc = False
    if doc is None:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(11)
        save_doc = True

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_data = []
    in_table = False
    in_code_block = False

    for line in lines:
        raw = line
        line = line.strip()

        # THOÁT KHỎI TABLE
        if in_table and not line.startswith('|'):
            add_table_to_doc(doc, table_data)
            table_data = []
            in_table = False
            if not line: # Nếu là dòng trống ngay sau table, có thể skip
                continue

        # CODE BLOCK
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(raw.rstrip('\n'))
            run.font.name = 'Courier New'
            continue

        # HEADINGS
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)

        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)

        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)

        # TABLE
        elif line.startswith('|'):
            if '---' in line:
                continue
            in_table = True
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(row)

        # IMAGE
        elif line.startswith('!['):
            handle_image(doc, line, md_path)

        # LIST
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_format(p, line[2:])

        elif re.match(r'\d+\. ', line):
            # Cắt bỏ phần "1. ", "2. "... ở đầu để tránh bị lặp số do style tự sinh
            prefix = re.match(r'\d+\. ', line).group()
            p = doc.add_paragraph(style='List Number')
            parse_inline_format(p, line[len(prefix):])

        # NORMAL TEXT
        else:
            if line:
                p = doc.add_paragraph()
                parse_inline_format(p, line)

    # Đảm bảo in cái table cuối cùng nếu file md kết thúc bằng table
    if table_data:
        add_table_to_doc(doc, table_data)

    if save_doc and docx_path:
        doc.save(docx_path)
        print(f"✅ Converted: {docx_path}")


# ==============================
# MAIN RUN
# ==============================
if __name__ == "__main__":

    md_files = [
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/tongquan.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichthuthap.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichxuly.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichdanhmuc.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichdulieumo.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichdulieuchu.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichcungcap.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichdieutiet.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichthongbao.md",
        "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/phantichquantri.md"
    ]

    # KHỞI TẠO FILE TỔNG HỢP (Chung 1 Document object để không bị mất ảnh khi gộp)
    merged_doc = Document()
    style = merged_doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(11)

    for i, md in enumerate(md_files):
        if not os.path.exists(md):
            print(f"❌ Không tìm thấy file: {md}")
            continue

        # 1. Tạo file docx riêng lẻ (như cũ)
        docx = md.replace(".md", ".docx")
        convert_md_to_docx(md, docx_path=docx)
        
        # 2. Sinh trực tiếp vào Document của file gộp tổng (thay vì nối từng dòng xml)
        convert_md_to_docx(md, doc=merged_doc)

        # Thêm sang trang giữa các file
        if i < len(md_files) - 1:
            merged_doc.add_page_break()

    # Lưu file gộp
    tong_hop_path = "D:/Tư pháp/KhoDLDC/DLDC_1/tailieu/tailieuphantich/TONG_HOP.docx"
    merged_doc.save(tong_hop_path)
    print(f"✅ Merged file: {tong_hop_path}")