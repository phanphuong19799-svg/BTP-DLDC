
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_data = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=4)
        
        # Tables
        elif line.startswith('|'):
            if '---' in line:
                continue
            in_table = True
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                table_data.append(row)
        
        # Images (Basic placeholder handling)
        elif line.startswith('![') or '*(Chèn hình ảnh' in line:
            if table_data:
                add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[IMAGE PLACEHOLDER: {line}]")
            run.italic = True
            
        # Figure captions
        elif '<p align="center"' in line:
            caption = re.sub('<[^<]+?>', '', line).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.bold = True
            run.font.size = Pt(10)

        # Normal text
        else:
            if in_table:
                if not line:
                    add_table_to_doc(doc, table_data)
                    table_data = []
                    in_table = False
                continue
            
            if line:
                doc.add_paragraph(line)

    if table_data:
        add_table_to_doc(doc, table_data)

    doc.save(docx_path)
    print(f"Success: {docx_path} created.")

def add_table_to_doc(doc, table_data):
    if not table_data:
        return
    
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

if __name__ == "__main__":
    src_md = "tailieu/tailieuphantich/phantichthuthap.md"
    out_docx = "tailieu/tailieuphantich/phantichthuthap.docx"
    convert_md_to_docx(src_md, out_docx)
